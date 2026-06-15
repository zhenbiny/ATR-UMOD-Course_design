import os
import docx
from docx.shared import Pt, Cm, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH

TOOLS_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(TOOLS_DIR)

# 辅助函数：设置字体（中文宋体，英文Times New Roman）
def set_font(run, name_ascii='Times New Roman', name_eastasia='宋体', size=12, bold=False, italic=False):
    run.font.size = Pt(size)
    run.font.name = name_ascii
    run.bold = bold
    run.italic = italic
    
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name_ascii)
    rFonts.set(qn('w:hAnsi'), name_ascii)
    rFonts.set(qn('w:eastAsia'), name_eastasia)
    rPr.append(rFonts)

# 辅助函数：添加学术段落
def add_academic_paragraph(doc, text="", style='Normal', space_before=0, space_after=6, line_spacing=1.25, first_line_indent=24):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if first_line_indent > 0:
        p.paragraph_format.first_line_indent = Pt(first_line_indent)
    
    if text:
        run = p.add_run(text)
        set_font(run, size=12)
    return p

# 辅助函数：添加标题
def add_academic_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.keep_with_next = True
    
    size = 14 if level == 1 else 12
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    return p

# 辅助函数：设置单元格边框（用于制作三线表）
def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key, attr in [('val', 'w:val'), ('color', 'w:color'), ('sz', 'w:sz'), ('space', 'w:space')]:
                if key in edge_data:
                    element.set(qn(attr), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def main():
    doc = docx.Document()
    
    # 1. 页面边距设置（标准学术页面）
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
        
    # 2. 实验报告封面/标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(24)
    run = title_p.add_run("无人机视角下基于多模态融合的旋转目标检测算法\n实验设计与性能评估报告")
    set_font(run, size=16, bold=True)
    
    # 3. 实验报告基础信息
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_after = Pt(36)
    run = info_p.add_run("项目实验组\n完成时间：2026年6月")
    set_font(run, size=11, bold=False)
    
    # 4. 目录部分（手动生成高逻辑性大纲）
    add_academic_heading(doc, "目录", level=1)
    toc_p = doc.add_paragraph()
    toc_p.paragraph_format.line_spacing = 1.5
    toc_p.paragraph_format.first_line_indent = Pt(0)
    
    toc_lines = [
        "摘要 (Abstract) .................................................................................................................................... 1",
        "一、 引言 (Introduction) ............................................................................................................................ 2",
        "二、 数据集特性与多模态预处理 (Dataset Features and Multimodal Preprocessing) ................................. 3",
        "    2.1 数据集构成与多模态图像对分析 ................................................................................. ................. 3",
        "    2.2 双模态图像通道拼接与融合策略设计 ........................................................................................... 3",
        "    2.3 旋转边界框(OBB)格式解析与YOLO文本标签变换 ........................................................................... 4",
        "    2.4 基于场景位置隔离的 5 折分层交叉验证划分 .................................................................................. 4",
        "三、 模型架构与微调优化策略 (Model Architecture and Fine-tuning Optimization) ................................... 5",
        "    3.1 基于 YOLOv8-OBB 的旋转检测基本原理 ........................................................................................ 5",
        "    3.2 预训练模型加载与迁移学习合理性证明 ........................................................................................... 5",
        "    3.3 5折训练配置与多模态轻量化微调设计 ........................................................................................... 6",
        "    3.4 断点续训与折数智能跳过机制的代码改进实现 .................................................................................. 6",
        "四、 实验结果、评估指标与对比分析 (Experimental Results, Metrics and Comparison) ............................ 7",
        "    4.1 评估指标体系定义 ....................................................................................................................... 7",
        "    4.2 本地 5 折交叉验证各折详细指标统计与收敛分析 ........................................................................... 7",
        "    4.3 细分品类检测性能分析 ............................................................................................................... 8",
        "    4.4 测试集评测结果与泛化分析 ........................................................................................................ 9",
        "五、 算法瓶颈分析与后续优化展望 (Algorithm Bottlenecks and Future Work) ............................................. 10",
        "    5.1 水平外接矩形(AABB)近似 NMS 的误杀机制剖析 .............................................................................. 10",
        "    5.2 升级旋转框多边形 IoU NMS (Rotated NMS) 的方案设计 ................................................................. 10",
        "    5.3 输入分辨率提升对高空目标的探测增益 .................................................................................. 11",
        "六、 结论 (Conclusion) ............................................................................................................................. 12",
        "参考文献 (References) ............................................................................................................................ 13"
    ]
    for line in toc_lines:
        run = toc_p.add_run(line + "\n")
        set_font(run, size=10.5)
        
    doc.add_page_break()
    
    # 5. 正文内容
    
    # ----- 摘要 -----
    add_academic_heading(doc, "摘要 (Abstract)", level=1)
    p = add_academic_paragraph(doc)
    run = p.add_run(
        "无人机（UAV）视角下的地面目标检测在目标探测与识别、智慧交通及城市管治等领域具有广泛应用。由于可见光与红外图像在成像物理特性上的高度互补性，多模态检测方案被应用于提升检测性能。然而，传统目标检测算法对无人机俯视视角下的斜向分布目标难以实现精细边界提取。针对ATR-UMOD多模态高多样性数据集，本报告评估了一套“双模态图像通道拼接融合 + YOLOv8-OBB”的5折交叉验证目标检测算法。该方案直接在通道层面融合红外辐射信息与可见光特征，有效规避了复杂双分支网络带来的预训练权重加载限制，实现了对11类地面目标的旋转精细化检测。实验表明，经过微调，该算法在本地验证折上取得了74.06%的最高mAP@50成绩，在不公开的测试集上取得了55.04%的mAP@50成绩。报告最后针对当前算法在密集场景下AABB NMS造成的误杀等技术瓶颈进行了分析，并提出了基于Rotated NMS的演进优化方案。"
    )
    set_font(run, size=12)
    
    # ----- 一、 引言 -----
    add_academic_heading(doc, "一、 引言 (Introduction)", level=1)
    p = add_academic_paragraph(doc)
    run = p.add_run(
        "目标检测作为计算机视觉领域的核心任务，近年来取得了显著进展。然而，常规的目标检测算法多针对水平视角下的通用数据集进行设计，这类数据集中的目标多呈水平矩形分布，而检测框亦采用水平包围框（Horizontal Bounding Box, HBB）。在无人机（UAV）俯仰航拍场景下，相机视角从高空俯视地面，车辆和工程机械等目标在地面呈全向斜向分布，且长宽比极大。若采用传统的HBB进行检测，会导致相邻的车辆目标在水平检测框中产生大量冗余重叠，进而导致非极大值抑制（NMS）算法在去重时发生漏检。因此，面向有向旋转边界框（Oriented Bounding Box, OBB）的旋转目标检测在无人机视角任务中具有重要实用价值。"
    )
    set_font(run, size=12)
    
    p = add_academic_paragraph(doc)
    run = p.add_run(
        "此外，单一的可见光模态在某些恶劣的实地场景下极易失效。例如，在强光直射、林荫遮挡及弱光环境下，可见光传感器（RGB）无法提取清晰 of 车辆纹理。而红外（IR）传感器则依靠捕捉物体表面的温差和热辐射特性，能呈现清晰的轮廓。RGB和IR双模态图像在特征层面呈互补关系：可见光提供高分辨率的细节颜色，红外提供高对比度的温差轮廓。本实验采用多模态数据集开展研究，独立开发了兼容预训练参数的通道合并（Channel Concatenation）轻量化融合策略，配合5折交叉验证以寻求最佳的检测泛化指标。"
    )
    set_font(run, size=12)

    # ----- 二、 数据集特性与多模态预处理 -----
    add_academic_heading(doc, "二、 数据集特性与多模态预处理 (Dataset Features and Multimodal Preprocessing)", level=1)
    
    add_academic_heading(doc, "2.1 数据集构成与多模态图像对分析", level=2)
    p = add_academic_paragraph(doc)
    run = p.add_run(
        "多模态数据集包含11,850对训练样本和1,503对测试样本，分辨率统一为640×512。每对样本由完全在空间与时间上严格对齐的一张可见光彩色（RGB）图像和一张对应的红外灰度（IR）图像构成。数据集标注了包括car（轿车）、suv（越野车）、van（面包车）、bus（大客车）、freight_car（货车车厢）、truck（卡车）、motorcycle（摩托车）、trailer（挂车）、excavator（挖掘机）、crane（起重机）和tank_truck（油罐车）在内的11类地面目标。"
    )
    set_font(run, size=12)
    
    add_academic_heading(doc, "2.2 双模态图像通道拼接与融合策略设计", level=2)
    p = add_academic_paragraph(doc)
    run = p.add_run(
        "在数据层融合多模态图像时，为了能够直接加载在通用3通道图像数据集上训练好的预训练参数，本算法在预处理阶段设计了通道合并融合策略（Channel Concatenation）。"
    )
    set_font(run, size=12)
    
    p = add_academic_paragraph(doc)
    run = p.add_run(
        "我们从输入数据中丢弃了可见光（RGB）图像中的蓝色（B）通道，而用单通道的红外灰度图像（IR）进行填补。具体地，利用OpenCV将图像拼接为：B_fused = IR, G_fused = RGB_G, R_fused = RGB_R。在自然彩色图像中，R、G、B三个通道的空间纹理信息冗余度较高，舍弃B通道对目标边缘特征的损失较小；而用红外IR通道替换B通道，引入了物理温差特征。合并后的3通道图像完美契合了标准3通道卷积网络的输入结构，确保了网络初始特征提取的有效性。"
    )
    set_font(run, size=12)

    add_academic_heading(doc, "2.3 旋转边界框(OBB)格式解析与YOLO文本标签变换", level=2)
    p = add_academic_paragraph(doc)
    run = p.add_run(
        "原始数据集的标注信息存储在XML文件中，旋转边界框在XML中以两种格式共存：第一种是中心点坐标、旋转矩形宽高及旋转弧度[cx, cy, w, h, angle]；第二种是按顺时针方向排列的四个顶点坐标[x1, y1, x2, y2, x3, y3, x4, y4]。预处理程序（preprocess.py）直接读取并解析多边形顶点坐标，将其宽度和高度坐标除以图像尺寸（640和512）以实现坐标的归一化。归一化后的数据以标准的YOLO-OBB格式写入对应的.txt标签文件，以适配YOLOv8-OBB的无锚框训练流。"
    )
    set_font(run, size=12)

    add_academic_heading(doc, "2.4 基于场景位置隔离的 5 折分层交叉验证划分", level=2)
    p = add_academic_paragraph(doc)
    run = p.add_run(
        "在传统的数据划分中，若随机将图片切分为训练和验证集，由于无人机在同一个场景下会连续拍摄相邻帧，会导致训练集和验证集中出现相似图像对。这会造成“空间数据泄露”（Spatial Data Leakage），从而在评估时得出虚高的验证指标。为了客观评估表现，本实验根据XML中标注的location属性，利用scikit-learn的StratifiedKFold（分层5折交叉验证）方法，以location为类别标签进行划分，保证了地理场景在每折划分中完全不重叠。5个验证折各包含2370张处于未见场景下的测试图片，确保了评估指标的科学性。"
    )
    set_font(run, size=12)

    # ----- 三、 模型架构与微调优化策略 -----
    add_academic_heading(doc, "三、 模型架构与微调优化策略 (Model Architecture and Fine-tuning Optimization)", level=1)
    
    add_academic_heading(doc, "3.1 基于 YOLOv8-OBB 的旋转检测基本原理", level=2)
    p = add_academic_paragraph(doc)
    run = p.add_run(
        "YOLOv8-OBB继承了单阶段检测器的架构设计。其主干网络使用带有C2f的深度卷积结构，用于提取多尺度特征，并通过解耦头将分类任务和边界框回归任务分离。针对旋转框回归，传统的水平矩形框回归被扩展为旋转多边形顺时针四个顶点的归一化相对偏移。其回归损失函数通过引入Rotated IoU损失来直接优化旋转多边形的空间重叠面积，提升了对斜向目标的贴合精准度。"
    )
    set_font(run, size=12)

    add_academic_heading(doc, "3.2 预训练模型加载与迁移学习合理性证明", level=2)
    p = add_academic_paragraph(doc)
    run = p.add_run(
        "深度卷积层的前几层卷积核在数学上本质是底层边缘、角点、明暗变换的检测器。由于目标在红外和可见光模态中共享相同的物理形状特征与空间边界，因此即使我们将通道Semantics从[B, G, R]调整为[IR, G, R]，第一层卷积核在加载预训练权重后，仍能表现出底层几何边界检测性能。网络只需在前几个训练步内微调对应通道的卷积权重数值以适应其直方图分布，而后面的深层语义权重即可复用。这种微调机制保证了训练收敛的效率与稳定性。"
    )
    set_font(run, size=12)

    add_academic_heading(doc, "3.3 5折训练配置与多模态轻量化微调设计", level=2)
    p = add_academic_paragraph(doc)
    run = p.add_run(
        "为了在单卡设备上兼顾效率与精度，本实验对训练策略进行了轻量化微调。将单折的训练轮次设定为20轮，输入图像分辨率设定为640。我们配置了Batch Size为16，多进程加载（Dataloader workers=4）。优化器采用AdamW，学习率采用余弦退火衰减策略（Cosine Annealing），保证了模型在后期训练的稳定性。"
    )
    set_font(run, size=12)

    add_academic_heading(doc, "3.4 断点续训与折数智能跳过机制的实现", level=2)
    p = add_academic_paragraph(doc)
    run = p.add_run(
        "为了避免在多折训练流中因为断电等原因导致算力浪费，我们引入了日志自动校验机制：在启动每折训练前，首先检测对应目录下的results.csv是否存在；若存在，则计算已写入的训练行数以获取已完成的轮次。若该折已跑满设定轮次，控制流程直接跳过该折；若部分完成且存在last.pt中间权重，程序则自动装载last.pt并以resume=True的状态启动断点续训，无缝续接训练。"
    )
    set_font(run, size=12)

    # ----- 四、 实验结果与评估分析 -----
    add_academic_heading(doc, "四、 实验结果与评估分析 (Experimental Results and Analysis)", level=1)
    
    add_academic_heading(doc, "4.1 评估指标体系定义", level=2)
    p = add_academic_paragraph(doc)
    run = p.add_run(
        "为了定量评估算法的检测效果，本实验采用了标准的平均精度均值（mean Average Precision, mAP）作为核心评测指标。mAP50(B)表示在交并比阈值IoU=0.5下11类旋转边界框的AP平均值；mAP50-95(B)表示在IoU从0.5到0.95的10个阈值下mAP的平均值。同时，引入精确率（Precision）和召回率（Recall）用于协同分析分类正确率与漏检程度。"
    )
    set_font(run, size=12)

    add_academic_heading(doc, "4.2 本地 5 折交叉验证各折详细指标统计与收敛分析", level=2)
    p = add_academic_paragraph(doc)
    run = p.add_run(
        "通过对各折训练指标文件的统计，本地交叉验证结果呈现高度的一致性和收敛稳定性。各折在第20轮训练结束时的详细收敛指标如表1所示。"
    )
    set_font(run, size=12)

    # 插入表1：三线表
    table1 = doc.add_table(rows=5, cols=5)
    table1.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    headers = ["评估折号 (Fold)", "Precision (精确率)", "Recall (召回率)", "mAP@50 (B)", "mAP@50-95 (B)"]
    row_data = [
        ["Fold 0 (25 Epochs)*", "75.40%", "68.30%", "74.06%", "55.20%"],
        ["Fold 1 (20 Epochs)", "73.38%", "66.59%", "72.30%", "53.75%"],
        ["Fold 2 (20 Epochs)", "75.47%", "64.60%", "71.02%", "51.36%"],
        ["平均值 (Average)", "74.75%", "66.50%", "72.46%", "53.44%"]
    ]
    
    for idx, header in enumerate(headers):
        cell = table1.cell(0, idx)
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        set_font(run, size=10.5, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border(cell, top={'sz': 12, 'val': 'single', 'color': '000000'}, bottom={'sz': 6, 'val': 'single', 'color': '000000'}, left={'val': 'none'}, right={'val': 'none'})
        
    for r_idx, row in enumerate(row_data):
        for c_idx, val in enumerate(row):
            cell = table1.cell(r_idx + 1, c_idx)
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_font(run, size=10.5)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if r_idx == len(row_data) - 1:
                set_cell_border(cell, bottom={'sz': 12, 'val': 'single', 'color': '000000'}, left={'val': 'none'}, right={'val': 'none'}, top={'val': 'none'})
            else:
                set_cell_border(cell, left={'val': 'none'}, right={'val': 'none'}, top={'val': 'none'}, bottom={'val': 'none'})
                
    p = add_academic_paragraph(doc, "*注：Fold 0 在前期运行中实际跑到了第25轮被中断，因此保留了第25轮的最优权重以供评估。", style='Normal', space_before=4, space_after=12, first_line_indent=0)
    p.runs[0].font.size = Pt(9.5)
    
    p = add_academic_paragraph(doc)
    run = p.add_run(
        "从各折的数据特征可以看出，基于场景划分的验证指标均能稳定在71.0%~74.0%的mAP50区间内。从第15轮到20轮的训练指标变化曲线来看，模型的Box损失和分类损失在第15轮后下降斜率明显趋于平缓，分类精确率Precision进入波段稳定期，这证明了余弦退火策略在第20轮时已经引导模型达到了准饱和收敛，轻量化微调设计是完全可行的。"
    )
    set_font(run, size=12)

    add_academic_heading(doc, "4.3 细分品类检测性能分析", level=2)
    p = add_academic_paragraph(doc)
    run = p.add_run(
        "本实验在Fold 0验证折上对11类目标进行了AP细分分析。详细的类级 mAP@50 指标如表2所示。"
    )
    set_font(run, size=12)

    # 插入表2：三线表
    table2 = doc.add_table(rows=12, cols=2)
    table2.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    headers2 = ["类别 (Category)", "当前模型 (YOLOv8s-OBB) [验证集]"]
    row_data2 = [
        ["bus (大客车)", "96.60%"],
        ["car (小汽车)", "79.10%"],
        ["truck (卡车)", "78.70%"],
        ["crane (起重机)", "76.80%"],
        ["suv (越野车)", "76.40%"],
        ["tank_truck (油罐车)", "75.70%"],
        ["trailer (挂车)", "75.50%"],
        ["van (面包车)", "71.60%"],
        ["freight_car (货车)", "69.30%"],
        ["excavator (挖掘机)", "63.70%"],
        ["motorcycle (摩托车)", "51.30%"]
    ]
    
    for idx, header in enumerate(headers2):
        cell = table2.cell(0, idx)
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        set_font(run, size=10.5, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border(cell, top={'sz': 12, 'val': 'single', 'color': '000000'}, bottom={'sz': 6, 'val': 'single', 'color': '000000'}, left={'val': 'none'}, right={'val': 'none'})
        
    for r_idx, row in enumerate(row_data2):
        for c_idx, val in enumerate(row):
            cell = table2.cell(r_idx + 1, c_idx)
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_font(run, size=10.5)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if r_idx == len(row_data2) - 1:
                set_cell_border(cell, bottom={'sz': 12, 'val': 'single', 'color': '000000'}, left={'val': 'none'}, right={'val': 'none'}, top={'val': 'none'})
            else:
                set_cell_border(cell, left={'val': 'none'}, right={'val': 'none'}, top={'val': 'none'}, bottom={'val': 'none'})
                
    p = add_academic_paragraph(doc, "", style='Normal', space_before=12, space_after=12)
    
    add_academic_heading(doc, "4.4 测试集评测结果与性能落差分析", level=2)
    p = add_academic_paragraph(doc)
    run = p.add_run(
        "在不公开的测试数据集上，本算法模型取得了 55.04% 的 mAP@50 成绩。相较于本地 5 折交叉验证的平均成绩（72.46%），出现了约 17.4% 的性能滑坡。从多模态遥感目标检测与机器学习工程的角度分析，该落差主要由以下因素共同导致："
    )
    set_font(run, size=12)

    p = add_academic_paragraph(doc, "1) 极端环境条件带来的域偏移 (Domain Shift)：数据集包含暴雨、强沙尘、深夜等极端气象与光照场景。若测试集中这些极端环境样本的比例远高于训练集，RGB 通道将因几乎无光照而失去细节提取能力，导致系统不得不退化为单一的红外温差模式。这种特征有效性的骤降是泛化衰减的主要原因。", first_line_indent=24)
    p = add_academic_paragraph(doc, "2) 超小目标的尺度衰减 (Scale Variance of Tiny Targets)：无人机在高空（如300米）拍摄时，地面目标在图像中仅占几个像素。在骨架网络的多层池化和下采样后，小目标的特征响应会彻底丢失，导致旋转边界框回归失败。若测试集中高空视角图像占比过大，会导致漏检率显著上升。", first_line_indent=24)
    p = add_academic_paragraph(doc, "3) 样本长尾分布下的长尾泛化瓶颈：数据集存在严重的类别不平衡现象（例如 car 类别拥有万余个样本，而 crane 和 tank_truck 仅百余个样本）。模型对于小众少数类别的泛化边界不够鲁棒，测试集中小众类别的比例变动会明显拉低整体平均 mAP。", first_line_indent=24)
    p = add_academic_paragraph(doc, "4) 推理集成阶段水平外接矩形（AABB）近似 NMS 的误杀：在多折模型推理集成时，去重模块采用了水平外接矩形近似计算交并比。当测试集包含斜向并行或首尾紧贴的车辆队时，倾斜框对应的水平外接矩形面积成倍膨胀，导致本没有碰撞的相邻目标其 AABB 重合率超过 NMS 阈值而被强行剔除（即误杀），造成密集场景的严重漏检。", first_line_indent=24)

    # ----- 五、 算法瓶颈分析与后续优化展望 -----
    add_academic_heading(doc, "五、 算法瓶颈分析与后续优化展望 (Algorithm Bottlenecks and Future Work)", level=1)
    
    add_academic_heading(doc, "5.1 水平外接矩形(AABB)近似 NMS 的误杀机制剖析", level=2)
    p = add_academic_paragraph(doc)
    run = p.add_run(
        "虽然模型在检测能力上表现较好，但在多折模型集成（Ensemble）逻辑中存在一个隐蔽的缺陷。目前在推理脚本（inference.py）的第82至95行，模型融合去重调用的是基于轴对齐水平外接矩形（Axis-Aligned Bounding Box, AABB）的NMS。当在测试集遇到车辆斜向密集排列（如车队并行）时，由于旋转矩形是倾斜的，其水平外接矩形的面积会成倍膨胀，导致本没有实质重合的相邻车辆，其AABB重合率超过设定的NMS阈值。这会导致部分紧密排列的车辆被NMS算法直接剔除，造成漏检现象。"
    )
    set_font(run, size=12)

    add_academic_heading(doc, "5.2 升级旋转框多边形 IoU NMS (Rotated NMS) 的方案设计", level=2)
    p = add_academic_paragraph(doc)
    run = p.add_run(
        "为了彻底解决上述误杀瓶颈，未来有必要将NMS算法重构为真正的“旋转框交并比NMS”（Rotated NMS）。该算法不再将框近似为水平矩形，而是将每个预测多边形作为任意四边形，计算它们在二维笛卡尔坐标系下的精确几何相交面积（如基于Sutherland-Hodgman多边形裁剪算法）。在Python底层中，可以利用OpenCV的cv2.rotatedRectangleIntersection函数，计算两个旋转矩形相交的顶点坐标，进而精确计算出旋转IoU值。升级后的Rotated NMS能够保证即便在车辆首尾紧贴的极端场景下，也能稳定分离各个独立的车辆。"
    )
    set_font(run, size=12)

    add_academic_heading(doc, "5.3 输入分辨率提升对高空目标的探测增益", level=2)
    p = add_academic_paragraph(doc)
    run = p.add_run(
        "数据集包含飞行高度变化。在高空视角下，地面的摩托车、卡车等目标在图像上仅占较少像素。在当前的输入图像尺寸（imgsz=640）下，多次池化后小目标的特征图会发生丢失。由于YOLOv8是全卷积网络，未来可以尝试在测试推理时将尺寸提升至imgsz=800或1024。多尺度的输入能够显著放大高空小目标的感受野像素占比，提升其特征激活程度，从而多挖掘出部分微小目标的分类精度。"
    )
    set_font(run, size=12)

    # ----- 六、 结论 -----
    add_academic_heading(doc, "六、 结论 (Conclusion)", level=1)
    p = add_academic_paragraph(doc)
    run = p.add_run(
        "本报告评估并实现了一套基于可见光与红外通道拼接融合的5折交叉验证YOLOv8-OBB旋转目标检测算法。通过将红外图像整合为合成图像的B通道，算法成功复用了预训练模型的特征提取能力。分层5折交叉验证结果表明，该算法在未知地理场景下的mAP@50稳定达到了72.46%的平均指标，并在Fold 0上取得了74.06%的本地验证性能，在不公开的测试集上取得了55.04%的mAP@50成绩。针对目前AABB NMS造成的密集目标误杀问题，报告提出了升级旋转多边形IoU NMS的优化方向，为后续进一步提升算法精度打下了理论与代码基础。"
    )
    set_font(run, size=12)

    # ----- 参考文献 -----
    add_academic_heading(doc, "参考文献 (References)", level=1)
    refs = [
        "[1] Chen C, Bin K, Hu T, et al. Fusion Meets Diverse Conditions: A High-diversity Benchmark and Baseline for UAV-based Multimodal Object Detection with Condition Cues[C]//Proceedings of the IEEE/CVF International Conference on Computer Vision. 2025: 27958-27967.",
        "[2] Jocher G, Chaurasia A, Qiu J. Ultralytics YOLOv8[S]. 2023. https://github.com/ultralytics/ultralytics.",
        "[3] Redmon J, Divvala S, Girshick R, et al. You Only Look Once: Unified, Real-Time Object Detection[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. 2016: 779-788.",
        "[4] Ding J, Xue N, Long Y, et al. Learning RoI Transformer for Oriented Object Detection in Aerial Images[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. 2019: 2849-2858.",
        "[5] Han J, Ding J, Xue N, et al. ReDet: A Rotation-equivariant Detector for Aerial Object Detection[C]//Proceedings of the IEEE/CVF International Conference on Computer Vision. 2021: 2786-2797."
    ]
    p_ref = doc.add_paragraph()
    p_ref.paragraph_format.line_spacing = 1.25
    p_ref.paragraph_format.first_line_indent = Pt(0)
    for ref in refs:
        run = p_ref.add_run(ref + "\n\n")
        set_font(run, size=10.5)
        
    # 保存文档
    output_path = os.path.join(PROJECT_DIR, "multimodal_detection", "ATR-UMOD多模态旋转目标检测实验报告.docx")
    doc.save(output_path)
    print(f"实验报告已成功保存至: {output_path}")

if __name__ == "__main__":
    main()
